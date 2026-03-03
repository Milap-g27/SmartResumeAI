"""Resume generator service — converts markdown resume to PDF and DOCX."""

from __future__ import annotations

import io
import os
import platform
import re
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    Table,
    TableStyle,
)
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ═══════════════════════════════════════════════════════════════════
# Unicode Font Registration (supports full UTF-8 in PDFs)
# ═══════════════════════════════════════════════════════════════════

_FONT = "Helvetica"
_FONT_BD = "Helvetica-Bold"
_FONT_IT = "Helvetica-Oblique"
_FONT_BI = "Helvetica-BoldOblique"

def _register_unicode_fonts():
    """Try to register a Unicode TTF font family; fall back to Helvetica."""
    global _FONT, _FONT_BD, _FONT_IT, _FONT_BI

    candidates: list[tuple[str, str, str, str]] = []

    if platform.system() == "Windows":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fd = os.path.join(windir, "Fonts")
        candidates.append((
            os.path.join(fd, "arial.ttf"),
            os.path.join(fd, "arialbd.ttf"),
            os.path.join(fd, "ariali.ttf"),
            os.path.join(fd, "arialbi.ttf"),
        ))
    else:
        # Linux (Render / Ubuntu / Debian)
        candidates += [
            (
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
            ),
            (
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf",
            ),
        ]

    for regular, bold, italic, bold_italic in candidates:
        if not os.path.isfile(regular):
            continue
        try:
            pdfmetrics.registerFont(TTFont("UniFont", regular))
            _FONT = "UniFont"
            _FONT_BD = "UniFont"
            _FONT_IT = "UniFont"
            _FONT_BI = "UniFont"

            if os.path.isfile(bold):
                pdfmetrics.registerFont(TTFont("UniFont-Bold", bold))
                _FONT_BD = "UniFont-Bold"
            if os.path.isfile(italic):
                pdfmetrics.registerFont(TTFont("UniFont-Italic", italic))
                _FONT_IT = "UniFont-Italic"
            if os.path.isfile(bold_italic):
                pdfmetrics.registerFont(TTFont("UniFont-BI", bold_italic))
                _FONT_BI = "UniFont-BI"

            pdfmetrics.registerFontFamily(
                "UniFont",
                normal=_FONT,
                bold=_FONT_BD,
                italic=_FONT_IT,
                boldItalic=_FONT_BI,
            )
            return
        except Exception:
            continue

_register_unicode_fonts()


# ═══════════════════════════════════════════════════════════════════
# PDF Generation (reportlab)
# ═══════════════════════════════════════════════════════════════════

def _build_pdf_styles():
    """Create ATS-safe PDF styles using Unicode-capable fonts."""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="ResumeName",
        parent=styles["Title"],
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=HexColor("#1a1a1a"),
        fontName=_FONT_BD,
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="ResumeContact",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=8,
        textColor=HexColor("#555555"),
        fontName=_FONT,
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="SectionHeading",
        parent=styles["Heading2"],
        fontSize=11,
        leading=14,
        spaceBefore=6,
        spaceAfter=2,
        textColor=HexColor("#1a1a1a"),
        fontName=_FONT_BD,
        borderWidth=0,
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="JobHeaderCompany",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        fontName=_FONT_BD,
        textColor=HexColor("#1a1a1a"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="JobHeaderLocation",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        alignment=TA_RIGHT,
        fontName=_FONT_BD,
        textColor=HexColor("#1a1a1a"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="JobHeaderRole",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        fontName=_FONT_IT,
        textColor=HexColor("#333333"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="JobHeaderDates",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        fontName=_FONT_IT,
        alignment=TA_RIGHT,
        textColor=HexColor("#333333"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="SubHeading",
        parent=styles["Heading3"],
        fontSize=10,
        leading=12,
        spaceBefore=6,
        spaceAfter=2,
        fontName=_FONT_BD,
        textColor=HexColor("#333333"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="ResumeBody",
        parent=styles["Normal"],
        fontSize=9.5,
        leading=11.5,
        spaceAfter=1,
        fontName=_FONT,
        textColor=HexColor("#333333"),
        wordWrap='CJK',
    ))

    styles.add(ParagraphStyle(
        name="BulletItem",
        parent=styles["Normal"],
        fontSize=9.5,
        leading=11.5,
        leftIndent=12,
        spaceAfter=1,
        bulletIndent=6,
        fontName=_FONT,
        textColor=HexColor("#333333"),
        wordWrap='CJK',
    ))

    return styles


def _sanitize_for_pdf(text: str) -> str:
    """Normalise special Unicode chars for consistent PDF rendering.

    With a registered TTF font all UTF-8 glyphs are supported, so we
    only replace invisible / ambiguous characters that could cause
    layout issues.  We intentionally keep accented, Cyrillic, CJK,
    Devanagari, Gujarati and other real characters intact.
    """
    replacements = {
        '\u2010': '-',    # hyphen
        '\u2011': '-',    # non-breaking hyphen
        '\u2012': '-',    # figure dash
        '\u2013': '-',    # en dash
        '\u2014': '--',   # em dash
        '\u2015': '--',   # horizontal bar
        '\u00ad': '-',    # soft hyphen
        '\u2018': "'",    # left single quote
        '\u2019': "'",    # right single quote
        '\u201a': "'",    # single low-9 quote
        '\u201c': '"',    # left double quote
        '\u201d': '"',    # right double quote
        '\u201e': '"',    # double low-9 quote
        '\u25a0': '-',    # black square
        '\u25cf': '-',    # black circle
        '\u25aa': '-',    # small black square
        '\u2022': '-',    # standard bullet
        '\u2026': '...',  # ellipsis
        '\u00a0': ' ',    # non-breaking space
        '\u200b': '',     # zero-width space
        '\u200e': '',     # LRM
        '\u200f': '',     # RLM
        '\u2028': '\n',   # line separator
        '\u2029': '\n',   # paragraph separator
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)

    # Escape XML-special characters for ReportLab Paragraph
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    return text


def generate_pdf(markdown_text: str) -> bytes:
    """Convert markdown resume text to ATS-safe PDF bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.35 * inch,
        rightMargin=0.35 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )
    styles = _build_pdf_styles()
    story = []

    lines = markdown_text.strip().split("\n")
    i = 0
    in_header = True
    while i < len(lines):
        line = lines[i].rstrip()

        # H1 — Name
        if line.startswith("# ") and not line.startswith("## "):
            text = _sanitize_for_pdf(line[2:].strip())
            story.append(Paragraph(text, styles["ResumeName"]))

        # H2 — Section heading
        elif line.startswith("## "):
            in_header = False
            raw = line[3:].strip().upper()
            text = _sanitize_for_pdf(raw)
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=HexColor("#cccccc"), spaceAfter=2, spaceBefore=6,
            ))
            story.append(Paragraph(f"<b>{text}</b>", styles["SectionHeading"]))

        # H3 — Sub heading (company/role)
        elif line.startswith("### "):
            text = line[4:].strip()
            if "|" in text:
                raw_parts = [p.strip() for p in text.split("|")]
                company = _sanitize_for_pdf(raw_parts[0].upper()) if len(raw_parts) > 0 else ""
                role = _sanitize_for_pdf(raw_parts[1]) if len(raw_parts) > 1 else ""
                location = _sanitize_for_pdf(raw_parts[2]) if len(raw_parts) > 2 else ""
                dates = _sanitize_for_pdf(raw_parts[3]) if len(raw_parts) > 3 else ""
                
                comp_para = Paragraph(f"<b>{company}</b>", styles["JobHeaderCompany"]) if company else Paragraph("", styles["JobHeaderCompany"])
                loc_para = Paragraph(f"<b>{location}</b>", styles["JobHeaderLocation"]) if location else Paragraph("", styles["JobHeaderLocation"])
                role_para = Paragraph(f"<i>{role}</i>", styles["JobHeaderRole"]) if role else Paragraph("", styles["JobHeaderRole"])
                date_para = Paragraph(f"<i>{dates}</i>", styles["JobHeaderDates"]) if dates else Paragraph("", styles["JobHeaderDates"])

                row1 = [comp_para, loc_para]
                row2 = [role_para, date_para]
                
                t = Table([row1, row2], colWidths=["70%", "30%"])
                t.setStyle(TableStyle([
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(t)
                story.append(Spacer(1, 4))
            else:
                text = _sanitize_for_pdf(text)
                story.append(Paragraph(f"<b>{text}</b>", styles["SubHeading"]))

        elif line.strip().startswith("- "):
            text = line.strip()[2:].strip()
            text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
            text = _sanitize_for_pdf(text)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"- {text}", styles["BulletItem"]))

        # Empty line
        elif not line.strip() or line.strip() in ['---', '***']:
            pass

        # Regular text
        else:
            text = line.strip()
            text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
            text = _sanitize_for_pdf(text)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            
            if in_header and "|" in text:
                story.append(Paragraph(text, styles["ResumeContact"]))
            elif text:
                story.append(Paragraph(text, styles["ResumeBody"]))

        i += 1

    doc.build(story)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
# DOCX Generation (python-docx)
# ═══════════════════════════════════════════════════════════════════

def generate_docx(markdown_text: str) -> bytes:
    """Convert markdown resume text to DOCX bytes."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    lines = markdown_text.strip().split("\n")

    for line in lines:
        line_stripped = line.rstrip()

        # H1 — Name
        if line_stripped.startswith("# ") and not line_stripped.startswith("## "):
            text = line_stripped[2:].strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(26, 26, 26)

        # H2 — Section heading
        elif line_stripped.startswith("## "):
            text = line_stripped[3:].strip()
            p = doc.add_heading(text.upper(), level=1)
            for run in p.runs:
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor(26, 26, 26)

        # H3 — Sub heading
        elif line_stripped.startswith("### "):
            text = line_stripped[4:].strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(51, 51, 51)

        # Bullet
        elif line_stripped.strip().startswith("- "):
            text = line_stripped.strip()[2:].strip()
            # Remove markdown bold markers
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)

        # Empty line — skip
        elif not line_stripped.strip():
            continue

        # Regular text
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line_stripped.strip())
            if text:
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
